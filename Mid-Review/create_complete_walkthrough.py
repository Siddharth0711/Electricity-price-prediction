"""
Create EXTREMELY Detailed Walkthrough Document
Complete step-by-step guide with explanations, usage, and deployment
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    return heading

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text, style=style)
    return p

def create_complete_walkthrough():
    """Generate complete detailed walkthrough"""
    doc = Document()
    
    # === TITLE PAGE ===
    title = doc.add_heading('COMPLETE STEP-BY-STEP WALKTHROUGH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Weather-Driven Electricity Price Forecasting System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.bold = True
    
    tagline = doc.add_paragraph('From Zero to Deployment: Every Step Explained in Detail')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('Created: February 13, 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Total Content: Comprehensive End-to-End Guide').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # === TABLE OF CONTENTS ===
    add_heading(doc, 'TABLE OF CONTENTS', 1)
    toc_items = [
        'PART 1: Understanding What We Built',
        'PART 2: Complete File Structure Explanation',
        'PART 3: Step-by-Step Component Guide (The Logic)',
        'PART 4: How to Run the Pipeline',
        'PART 5: GitHub Deployment Guide',
        'PART 6: Building the Streamlit Application',
        'PART 7: Troubleshooting and Best Practices'
    ]
    for item in toc_items:
        add_paragraph(doc, item, style='List Bullet')
    
    doc.add_page_break()
    
    # === PART 1: UNDERSTANDING THE SYSTEM ===
    add_heading(doc, 'PART 1: UNDERSTANDING WHAT WE BUILT', 1)
    
    add_heading(doc, '1.1 The Business Problem', 2)
    add_paragraph(doc, "In the electricity market, prices (Market Clearing Prices or MCP) are extremely volatile. Unlike most commodities, electricity cannot be easily stored in large quantities, so supply must match demand in real-time. This system models how weather conditions—like sunshine, wind speed, and rain—affect the generation of renewable energy, which in turn shifts the supply curve and changes the market price. By predicting these shifts, traders can optimize their bids and increase profitability.")
    
    add_heading(doc, '1.2 Multi-Source Energy Modeling', 2)
    add_paragraph(doc, "We don't just look at 'renewables' as a single block. We model:")
    add_paragraph(doc, "• Solar: Dependent on irradiance and temperature.", style='List Bullet')
    add_paragraph(doc, "• Wind: Dependent on wind speed and power curves.", style='List Bullet')
    add_paragraph(doc, "• Hydro: Dependent on rainfall patterns.", style='List Bullet')
    add_paragraph(doc, "• Non-Renewables: Coal, Gas, and Nuclear, which fill the gap based on their marginal costs.", style='List Bullet')
    
    doc.add_page_break()

    # === PART 2: FILE STRUCTURE ===
    add_heading(doc, 'PART 2: COMPLETE FILE STRUCTURE EXPLANATION', 1)
    add_paragraph(doc, "The project is organized to be modular and scalable. Here is a detailed breakdown of the folders and key files.")
    
    add_heading(doc, '2.1 Root Folder', 2)
    add_paragraph(doc, "• README.md: The project overview and entry point for documentation.", style='List Bullet')
    add_paragraph(doc, "• QUICKSTART.md: A fast-track guide to getting the system running.", style='List Bullet')
    add_paragraph(doc, "• requirements.txt: All Python libraries needed (pandas, scikit-learn, xgboost, etc.).", style='List Bullet')
    
    add_heading(doc, '2.2 config/ Directory', 2)
    add_paragraph(doc, "This folder contains YAML files where you can change system parameters without touching the code.")
    add_paragraph(doc, "• data_sources.yaml: Define plant capacities (MW) and weather dependencies.", style='List Bullet')
    add_paragraph(doc, "• model_config.yaml: Hyperparameters for ML models like XGBoost and LightGBM.", style='List Bullet')
    add_paragraph(doc, "• trading_config.yaml: Rules for bidding and signal thresholds.", style='List Bullet')
    
    add_heading(doc, '2.3 src/ Directory', 2)
    add_paragraph(doc, "This is where the 'brain' of the application lives.")
    add_paragraph(doc, "• data/: Scripts for collection and cleaning.", style='List Bullet')
    add_paragraph(doc, "• features/: Logic for physics-based renewable mapping and supply mix.", style='List Bullet')
    add_paragraph(doc, "• models/: Training and evaluation logic.", style='List Bullet')
    add_paragraph(doc, "• trading/: Signal generation and P&L tracking.", style='List Bullet')
    
    doc.add_page_break()

    # === PART 3: STEP-BY-STEP COMPONENT GUIDE ===
    add_heading(doc, 'PART 3: STEP-BY-STEP COMPONENT GUIDE (THE LOGIC)', 1)
    
    add_heading(doc, '3.1 Data Collection (data_collector.py)', 2)
    add_paragraph(doc, "Why we do it: To build a model, we need data. This script generates 2 years of hourly synthetic data that mimics real-world patterns like seasonal temperatures and peak-hour demand.")
    add_paragraph(doc, "Reasoning: Real market data is often pay-walled. Synthetic generation allows us to build and test the architecture perfectly before switching to live APIs.")
    
    add_heading(doc, '3.2 Weather to Renewable Mapping (weather_renewable_mapper.py)', 2)
    add_paragraph(doc, "Why we do it: We use physics equations to estimate MW output.")
    add_paragraph(doc, "• Solar: Efficiency × Irradiance × Temperature correction.", style='List Bullet')
    add_paragraph(doc, "• Wind: Follows a 'Power Curve' (0 output below cut-in speed, max output at rated speed).", style='List Bullet')
    
    add_heading(doc, '3.3 Supply Mix & Merit Order (supply_mix_calculator.py)', 2)
    add_paragraph(doc, "Why we do it: This is the most critical part. We calculate 'Shortfall = Demand - Renewables'. We then 'dispatch' non-renewables in order of their cost (Nuclear → Coal → Gas). The last plant needed sets the price (MCP).")
    
    add_heading(doc, '3.4 Model Training (model_trainer.py)', 2)
    add_paragraph(doc, "Why we do it: We train multiple models (XGBoost, LightGBM, Random Forest) because some perform better on price spikes, while others are better at base trends. We automatically select the best one.")
    
    doc.add_page_break()

    # === PART 4: HOW TO RUN THE PIPELINE ===
    add_heading(doc, 'PART 4: HOW TO RUN THE PIPELINE', 1)
    add_paragraph(doc, "Follow these steps in your terminal inside the FRESH START folder:")
    
    add_paragraph(doc, "Step 1: Setup Environment", style='List Number')
    add_paragraph(doc, "python -m venv venv\nsource venv/bin/activate  # On Mac\npip install -r requirements.txt")
    
    add_paragraph(doc, "Step 2: Run Full Pipeline", style='List Number')
    add_paragraph(doc, "python src/main.py --mode full")
    add_paragraph(doc, "Explanation: This one command runs data generation, feature building, training, and trading insights sequentially.")
    
    add_paragraph(doc, "Step 3: Run Specific Stages", style='List Number')
    add_paragraph(doc, "• Forecast only: python src/main.py --mode forecast --horizon 24", style='List Bullet')
    add_paragraph(doc, "• Train only: python src/main.py --mode train", style='List Bullet')
    
    doc.add_page_break()

    # === PART 5: GITHUB DEPLOYMENT GUIDE ===
    add_heading(doc, 'PART 5: GITHUB DEPLOYMENT GUIDE', 1)
    add_paragraph(doc, "To host your project on GitHub and collaborate:")
    
    add_heading(doc, '5.1 Initialize Repository', 2)
    add_paragraph(doc, "1. Create a new repo on GitHub.com.")
    add_paragraph(doc, "2. terminal: git init")
    add_paragraph(doc, "3. git add .")
    add_paragraph(doc, "4. git commit -m 'Initial project setup'")
    add_paragraph(doc, "5. git remote add origin <your-repo-url>")
    add_paragraph(doc, "6. git push -u origin main")
    
    add_heading(doc, '5.2 Setting Up GitHub Actions (CI/CD)', 2)
    add_paragraph(doc, "We added a .github/workflows/main.yml (optional but recommended) to run tests automatically whenever you push code. This ensures your changes didn't break the data pipeline.")
    
    doc.add_page_break()

    # === PART 6: BUILDING THE STREAMLIT APPLICATION ===
    add_heading(doc, 'PART 6: BUILDING THE STREAMLIT APPLICATION', 1)
    add_paragraph(doc, "Streamlit is the fastest way to turn your forecasting model into a web app.")
    
    add_heading(doc, '6.1 Creating steamlit_app.py', 2)
    add_paragraph(doc, "Write a new file named streamlit_app.py in the root folder with this logic:")
    add_paragraph(doc, "• Title: MCP Forecasting Dashboard", style='List Bullet')
    add_paragraph(doc, "• Sidebar: Let users adjust Capacity and Weather inputs.", style='List Bullet')
    add_paragraph(doc, "• Main: Show charts using st.line_chart() for forecasts.", style='List Bullet')
    add_paragraph(doc, "• Metrics: Display predicted P&L and signal (BUY/SELL).", style='List Bullet')
    
    add_paragraph(doc, "To run the app: streamlit run streamlit_app.py")
    
    doc.add_page_break()

    # === PART 7: TROUBLESHOOTING ===
    add_heading(doc, 'PART 7: TROUBLESHOOTING AND BEST PRACTICES', 1)
    add_paragraph(doc, "• Dependency Errors: If 'module not found', rerun pip install -r requirements.txt.", style='List Bullet')
    add_paragraph(doc, "• Data Path Issues: Always run scripts from the project root (FRESH START folder).", style='List Bullet')
    add_paragraph(doc, "• Model Not Training: Check if your synthetic data generation was successful (look in data/raw/).", style='List Bullet')
    
    add_heading(doc, 'Final Word', 2)
    add_paragraph(doc, "This system is now yours. You can customize the YAML files to reflect any energy market in the world.")

    return doc

if __name__ == "__main__":
    print("Generating the extremely detailed walkthrough document...")
    doc = create_complete_walkthrough()
    output_path = "Step_by_Step_Final_Walkthrough.docx"
    doc.save(output_path)
    print(f"✓ Document saved: {output_path}")
